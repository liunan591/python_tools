# -*- coding: utf-8 -*-

#%%
import docx
import getopt
import sys
import re

#%%
def get_photo_list(doc):
    list = []
    for para in doc.paragraphs:
        if len(para.text) < 50:
            if re.match(r"^\s*图(?=[\s\d])",para.text):
                list.append(para.text.strip())
    return list

def get_table_list(doc):
    list = []
    for para in doc.paragraphs:
        if len(para.text) < 50:
            if re.match(r"^\s*表(?=[\s\d])",para.text):
                list.append(para.text.strip())
    return list

def write_list_to_file(file_o,*args):
    with open(file_o,"w") as f:
        for arg in args:
            for a in arg:
                f.write(a)
                f.write("\n")

def main():
    opts,args = getopt.getopt(sys.argv[1:],"i:o:h")
    file_i = None
    file_o = None
    for o,a in opts:
        if o == "-i":
            file_i = a
        if o == "-o":
            file_o = a
        if o == "-h":
            print("获取docx文档中的图表目录\n调用参数：-i docx文件路径 -o 图表目录输出路径（xxx.txt）")
    if file_i and file_o:
        print("待处理文件为：%s"%file_i)
        print("输出文件为：%s"%file_o)
        doc = docx.Document(file_i)
        photo_list = get_photo_list(doc)
        table_list = get_table_list(doc)
        write_list_to_file(file_o,photo_list,table_list)
        
if __name__ == "__main__":
    main()
    